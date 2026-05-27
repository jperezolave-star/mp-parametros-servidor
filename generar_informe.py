import sys, json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def rgb(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color.lstrip('#'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph(doc, text, bold=False, size=10, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    return p

def add_cell_text(cell, text, bold=False, size=9, color="000000", align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text) if text is not None else "—")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)

def v(val):
    if val is None or val == "": return "—"
    return str(val)

def generar_informe(datos, output_path):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── PORTADA ──────────────────────────────────────────────────────────
    # Logo
    logo_path = "/app/logo_transparente.png"
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(10))

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("INFORME DE MANTENCIÓN ELÉCTRICA")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = rgb("1F4E79")

    # Edificio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datos.get("edificio","").upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = rgb("2E75B6")

    doc.add_paragraph()

    # Tabla de datos portada
    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = 'Table Grid'
    filas_datos = [
        ("Período", f"{datos.get('mes','')} {datos.get('anio','')}"),
        ("Técnico", datos.get("tecnico","")),
        ("Jefe de Mantención", datos.get("jefe_mantencion","—")),
        ("Fecha", datos.get("fecha","")),
        ("Tableros revisados", str(len(datos.get("tableros",[])))),
    ]
    for i, (lbl, val) in enumerate(filas_datos):
        c0, c1 = tabla.rows[i].cells
        set_cell_bg(c0, "1F4E79"); set_cell_border(c0, "2E75B6")
        set_cell_bg(c1, "EBF3FB"); set_cell_border(c1, "2E75B6")
        add_cell_text(c0, lbl, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)
        add_cell_text(c1, val, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # ── INTRODUCCIÓN ─────────────────────────────────────────────────────
    add_paragraph(doc, "Introducción", bold=True, size=16, color="1F4E79", space_before=6, space_after=4)
    add_paragraph(doc,
        f"El presente informe expone los resultados de la toma de parámetros eléctricos realizada en "
        f"{datos.get('edificio','')} durante el mes de {datos.get('mes','')} de {datos.get('anio','')}. "
        f"Este trabajo busca asegurar el correcto funcionamiento y seguridad de la red eléctrica.",
        size=10, space_after=8)

    add_paragraph(doc, "Trabajos ejecutados", bold=True, size=12, color="2E75B6", space_before=4, space_after=2)
    trabajos = [
        "Limpieza periódica mensual de tableros eléctricos.",
        "Limpieza periódica mensual de gabinete eléctrico con solvente dieléctrico.",
        "Inspección a componentes en tableros eléctricos.",
        "Toma de parámetros eléctricos: voltaje, corrientes, secuencias.",
        "Toma de temperatura de conductores y protecciones.",
    ]
    for t in trabajos:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(t)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ── PARÁMETROS ────────────────────────────────────────────────────────
    add_paragraph(doc, "Parámetros registrados", bold=True, size=16, color="1F4E79", space_before=6, space_after=6)

    tbs = datos.get("tableros", [])
    for tb in tbs:
        # Header del tablero
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        cell = t.rows[0].cells[0]
        set_cell_bg(cell, "1F4E79")
        set_cell_border(cell, "1F4E79")
        nombre = f"N° {tb.get('n','')}  —  {tb.get('nombre','')}   Piso {tb.get('piso','')}  |  {tb.get('servicio','')}"
        if tb.get('fecha_hora'):
            nombre += f"   ·   {tb.get('fecha_hora','')}"
        add_cell_text(cell, nombre, bold=True, color="FFFFFF", size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

        # Tabla de valores
        t2 = doc.add_table(rows=4, cols=10)
        t2.style = 'Table Grid'

        # Fila headers tensiones
        hdrs1 = ["R/S","R/T","S/T","R/N","S/N","T/N","CTE R","CTE S","CTE T","CTE N"]
        for j, h in enumerate(hdrs1):
            c = t2.rows[0].cells[j]
            bg = "DDEBF7" if j < 6 else "E2EFDA"
            set_cell_bg(c, bg)
            set_cell_border(c)
            add_cell_text(c, h, bold=True, size=8)

        vals1 = [tb.get('v_rs'),tb.get('v_rt'),tb.get('v_st'),tb.get('v_rn'),tb.get('v_sn'),tb.get('v_tn'),
                 tb.get('i_r'),tb.get('i_s'),tb.get('i_t'),tb.get('i_n')]
        for j, val in enumerate(vals1):
            c = t2.rows[1].cells[j]
            set_cell_border(c)
            add_cell_text(c, v(val), size=8)

        hdrs2 = ["CTE TIE","R/TIE","S/TIE","T/TIE","NEU/TIE","Tº R","Tº S","Tº T","Tº N","Tº TIE"]
        for j, h in enumerate(hdrs2):
            c = t2.rows[2].cells[j]
            bg = "E2EFDA" if j == 0 else ("DDEBF7" if j < 5 else "FFF2CC")
            set_cell_bg(c, bg)
            set_cell_border(c)
            add_cell_text(c, h, bold=True, size=8)

        vals2 = [tb.get('i_tie'),tb.get('v_rtie'),tb.get('v_stie'),tb.get('v_ttie'),tb.get('v_ntie'),
                 tb.get('t_r'),tb.get('t_s'),tb.get('t_t'),tb.get('t_n'),tb.get('t_tie')]
        for j, val in enumerate(vals2):
            c = t2.rows[3].cells[j]
            set_cell_border(c)
            add_cell_text(c, v(val), size=8)

        # Observaciones
        obs = tb.get('obs','')
        if obs:
            add_paragraph(doc, f"Obs: {obs}", size=9, color="4A4A4A", space_before=1, space_after=4)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ── GRÁFICOS ──────────────────────────────────────────────────────────
    graf_paths = datos.get("_graf_paths", {})
    if graf_paths:
        add_paragraph(doc, "Análisis histórico de parámetros", bold=True, size=16, color="1F4E79", space_before=6, space_after=4)
        titulos = {
            "tension":       "Tensiones históricas (V)",
            "corriente":     "Corrientes históricas (A)",
            "temperatura":   "Temperaturas históricas (°C)",
            "barras_tension":"Tensión R/S por tablero — mes actual",
        }
        for key, titulo in titulos.items():
            path = graf_paths.get(key)
            if path and os.path.exists(path):
                add_paragraph(doc, titulo, bold=True, size=11, color="2E75B6", space_before=4, space_after=2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(path, width=Cm(16))
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
        doc.add_page_break()

    # ── FOTOGRAFÍAS ───────────────────────────────────────────────────────
    fotos = datos.get("fotografias", [])
    add_paragraph(doc, "Registro fotográfico", bold=True, size=16, color="1F4E79", space_before=6, space_after=4)
    if fotos:
        import base64, tempfile as tf2
        for i in range(0, len(fotos), 2):
            fila = fotos[i:i+2]
            t3 = doc.add_table(rows=1, cols=len(fila))
            t3.style = 'Table Grid'
            for j, foto in enumerate(fila):
                cell = t3.rows[0].cells[j]
                set_cell_border(cell, "DDDDDD")
                try:
                    data = foto.get("data","")
                    if data:
                        img_bytes = base64.b64decode(data)
                        tmp_img = tf2.NamedTemporaryFile(suffix=".jpg", delete=False)
                        tmp_img.write(img_bytes); tmp_img.close()
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(tmp_img.name, width=Cm(7))
                        os.unlink(tmp_img.name)
                except: pass
                add_cell_text(cell, foto.get("tablero",""), bold=True, size=9)
                if foto.get("descripcion"):
                    add_cell_text(cell, foto.get("descripcion",""), size=8, color="666666")
                add_cell_text(cell, foto.get("fecha_hora",""), size=7, color="999999")
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
    else:
        add_paragraph(doc, "No se adjuntaron fotografías en esta toma de parámetros.", size=10, color="888888")

    doc.add_page_break()

    # ── CONCLUSIONES ──────────────────────────────────────────────────────
    add_paragraph(doc, "Observaciones y conclusiones", bold=True, size=16, color="1F4E79", space_before=6, space_after=4)
    con_obs = [tb for tb in tbs if tb.get('obs','').strip()]
    if con_obs:
        add_paragraph(doc, "Se registraron las siguientes observaciones:", size=10, space_after=4)
        for tb in con_obs:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"Tablero {tb.get('nombre','')} (Piso {tb.get('piso','')}): ")
            run.bold = True; run.font.size = Pt(10)
            run2 = p.add_run(tb.get('obs',''))
            run2.font.size = Pt(10)
    else:
        add_paragraph(doc, "No se registraron observaciones relevantes. Los parámetros se encuentran dentro de los rangos normales.", size=10)

    # Firma
    doc.add_paragraph()
    t4 = doc.add_table(rows=1, cols=2)
    c0, c1 = t4.rows[0].cells
    for c in [c0, c1]:
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)

    add_cell_text(c0, "Elaborado por:", bold=True, size=10, color="1F4E79", align=WD_ALIGN_PARAGRAPH.LEFT)
    c0.add_paragraph()
    add_cell_text(c0, "_______________________________", size=10, color="888888", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell_text(c0, datos.get("tecnico",""), bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell_text(c0, "Técnico Electricista", size=9, color="666666", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_cell_text(c1, "Revisado por:", bold=True, size=10, color="1F4E79", align=WD_ALIGN_PARAGRAPH.LEFT)
    c1.add_paragraph()
    add_cell_text(c1, "_______________________________", size=10, color="888888", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell_text(c1, datos.get("jefe_mantencion",""), bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell_text(c1, "Jefe de Mantención", size=9, color="666666", align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(output_path)
    print(f"Informe guardado en: {output_path}")

if __name__ == "__main__":
    json_path = sys.argv[1] if len(sys.argv) > 1 else None
    out_path  = sys.argv[2] if len(sys.argv) > 2 else "/tmp/informe_test.docx"
    if json_path:
        with open(json_path) as f:
            datos = json.load(f)
    else:
        datos = {"edificio":"Torre B","mes":"Mayo","anio":"2026","tecnico":"Test","jefe_mantencion":"Jefe","fecha":"27-05-2026","tableros":[]}
    generar_informe(datos, out_path)
